import json
import os
from dataclasses import dataclass, asdict
from datetime import datetime
from pathlib import Path
import tkinter as tk
from tkinter import ttk, messagebox, filedialog

APP_DIR = Path.home() / ".resume_chunk_builder"
DATA_FILE = APP_DIR / "resume_data.json"


@dataclass
class Internship:
    name: str
    company: str
    period: str
    duty: str


@dataclass
class Project:
    name: str
    company: str
    period: str
    duty: str


@dataclass
class Skill:
    text: str


@dataclass
class Evaluation:
    text: str


DEFAULT_DATA = {
    "education": "",
    "internships": [],
    "projects": [],
    "skills": [],
    "evaluations": [],
}


class Storage:
    def __init__(self, path: Path):
        self.path = path
        APP_DIR.mkdir(parents=True, exist_ok=True)

    def load(self):
        if not self.path.exists():
            return DEFAULT_DATA.copy()
        with open(self.path, "r", encoding="utf-8") as f:
            data = json.load(f)
        for k, v in DEFAULT_DATA.items():
            data.setdefault(k, v)
        return data

    def save(self, data):
        with open(self.path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)


class MultiItemEditor(ttk.LabelFrame):
    def __init__(self, master, title, fields, data_key, data_store):
        super().__init__(master, text=title, padding=10)
        self.fields = fields
        self.data_key = data_key
        self.data_store = data_store

        self.listbox = tk.Listbox(self, height=6)
        self.listbox.grid(row=0, column=0, rowspan=len(fields) + 1, padx=(0, 10), sticky="nsew")
        self.listbox.bind("<<ListboxSelect>>", self.on_select)

        self.inputs = {}
        for i, field in enumerate(fields):
            ttk.Label(self, text=field).grid(row=i, column=1, sticky="w")
            entry = ttk.Entry(self, width=40)
            entry.grid(row=i, column=2, padx=5, pady=2, sticky="ew")
            self.inputs[field] = entry

        btn_frame = ttk.Frame(self)
        btn_frame.grid(row=len(fields), column=1, columnspan=2, pady=(6, 0), sticky="w")

        ttk.Button(btn_frame, text="新增", command=self.add_item).grid(row=0, column=0, padx=4)
        ttk.Button(btn_frame, text="更新", command=self.update_item).grid(row=0, column=1, padx=4)
        ttk.Button(btn_frame, text="删除", command=self.delete_item).grid(row=0, column=2, padx=4)

        self.columnconfigure(2, weight=1)
        self.rowconfigure(0, weight=1)

    def _list_label(self, item):
        if "name" in item:
            return item.get("name", "(未命名)")
        return item.get("text", "(空)")

    def refresh(self):
        self.listbox.delete(0, tk.END)
        for item in self.data_store[self.data_key]:
            self.listbox.insert(tk.END, self._list_label(item))

    def get_form_data(self):
        result = {}
        for field in self.fields:
            result[self.field_map(field)] = self.inputs[field].get().strip()
        return result

    @staticmethod
    def field_map(field_name):
        mapping = {
            "实习名称": "name",
            "项目名称": "name",
            "实习单位": "company",
            "项目单位/团队": "company",
            "实习时间段": "period",
            "项目时间段": "period",
            "实习职责": "duty",
            "项目职责": "duty",
            "技能": "text",
            "评价": "text",
        }
        return mapping[field_name]

    def clear_inputs(self):
        for entry in self.inputs.values():
            entry.delete(0, tk.END)

    def add_item(self):
        item = self.get_form_data()
        if not any(item.values()):
            messagebox.showwarning("提示", "请至少填写一个字段")
            return
        self.data_store[self.data_key].append(item)
        self.refresh()
        self.clear_inputs()

    def update_item(self):
        sel = self.listbox.curselection()
        if not sel:
            messagebox.showwarning("提示", "请先选择要更新的条目")
            return
        idx = sel[0]
        self.data_store[self.data_key][idx] = self.get_form_data()
        self.refresh()

    def delete_item(self):
        sel = self.listbox.curselection()
        if not sel:
            messagebox.showwarning("提示", "请先选择要删除的条目")
            return
        del self.data_store[self.data_key][sel[0]]
        self.refresh()
        self.clear_inputs()

    def on_select(self, _event):
        sel = self.listbox.curselection()
        if not sel:
            return
        item = self.data_store[self.data_key][sel[0]]
        for field in self.fields:
            key = self.field_map(field)
            self.inputs[field].delete(0, tk.END)
            self.inputs[field].insert(0, item.get(key, ""))


class ResumeGeneratorApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("简历分块生成工具")
        self.geometry("1300x780")

        self.storage = Storage(DATA_FILE)
        self.data = self.storage.load()

        self._build_ui()
        self.load_selection_lists()

    def _build_ui(self):
        notebook = ttk.Notebook(self)
        notebook.pack(fill="both", expand=True, padx=10, pady=10)

        edit_tab = ttk.Frame(notebook)
        generate_tab = ttk.Frame(notebook)
        notebook.add(edit_tab, text="内容管理")
        notebook.add(generate_tab, text="按需生成")

        self._build_edit_tab(edit_tab)
        self._build_generate_tab(generate_tab)

    def _build_edit_tab(self, parent):
        parent.columnconfigure(0, weight=1)
        parent.columnconfigure(1, weight=1)

        edu_frame = ttk.LabelFrame(parent, text="教育经历（单文本框）", padding=10)
        edu_frame.grid(row=0, column=0, columnspan=2, sticky="nsew", pady=(0, 10))
        self.edu_text = tk.Text(edu_frame, height=5)
        self.edu_text.pack(fill="both", expand=True)
        self.edu_text.insert("1.0", self.data.get("education", ""))

        self.intern_editor = MultiItemEditor(
            parent,
            "实习经历（可新增多条）",
            ["实习名称", "实习单位", "实习时间段", "实习职责"],
            "internships",
            self.data,
        )
        self.intern_editor.grid(row=1, column=0, sticky="nsew", padx=(0, 8), pady=6)

        self.project_editor = MultiItemEditor(
            parent,
            "项目经历（可新增多条）",
            ["项目名称", "项目单位/团队", "项目时间段", "项目职责"],
            "projects",
            self.data,
        )
        self.project_editor.grid(row=1, column=1, sticky="nsew", padx=(8, 0), pady=6)

        self.skill_editor = MultiItemEditor(
            parent,
            "个人技能（可新增多条）",
            ["技能"],
            "skills",
            self.data,
        )
        self.skill_editor.grid(row=2, column=0, sticky="nsew", padx=(0, 8), pady=6)

        self.eval_editor = MultiItemEditor(
            parent,
            "个人评价（可新增多条）",
            ["评价"],
            "evaluations",
            self.data,
        )
        self.eval_editor.grid(row=2, column=1, sticky="nsew", padx=(8, 0), pady=6)

        for i in range(3):
            parent.rowconfigure(i, weight=1)

        self.intern_editor.refresh()
        self.project_editor.refresh()
        self.skill_editor.refresh()
        self.eval_editor.refresh()

        bottom = ttk.Frame(parent)
        bottom.grid(row=3, column=0, columnspan=2, sticky="ew", pady=8)
        ttk.Button(bottom, text="保存全部内容", command=self.save_all).pack(side="left")
        ttk.Button(bottom, text="刷新“按需生成”可选项", command=self.load_selection_lists).pack(side="left", padx=8)

    def _build_generate_tab(self, parent):
        parent.columnconfigure(0, weight=1)
        parent.columnconfigure(1, weight=1)

        ttk.Label(parent, text="岗位关键词（可选，用于写在生成文件头部）").grid(
            row=0, column=0, columnspan=2, sticky="w", pady=(4, 0)
        )
        self.job_entry = ttk.Entry(parent)
        self.job_entry.grid(row=1, column=0, columnspan=2, sticky="ew", pady=(0, 8))

        self.sel_boxes = {}
        self.sel_vars = {}

        self._build_selector(parent, "实习经历", "internships", row=2, col=0)
        self._build_selector(parent, "项目经历", "projects", row=2, col=1)
        self._build_selector(parent, "个人技能", "skills", row=3, col=0)
        self._build_selector(parent, "个人评价", "evaluations", row=3, col=1)

        edu_block = ttk.LabelFrame(parent, text="教育经历（默认全部纳入）", padding=10)
        edu_block.grid(row=4, column=0, columnspan=2, sticky="nsew", pady=8)
        self.preview_edu = tk.Text(edu_block, height=6)
        self.preview_edu.pack(fill="both", expand=True)

        action = ttk.Frame(parent)
        action.grid(row=5, column=0, columnspan=2, sticky="ew", pady=8)
        ttk.Button(action, text="导出 TXT", command=lambda: self.export("txt")).pack(side="left", padx=4)
        ttk.Button(action, text="导出 Markdown", command=lambda: self.export("md")).pack(side="left", padx=4)
        ttk.Button(action, text="导出 DOCX", command=lambda: self.export("docx")).pack(side="left", padx=4)
        ttk.Button(action, text="导出 PDF", command=lambda: self.export("pdf")).pack(side="left", padx=4)

        for i in range(6):
            parent.rowconfigure(i, weight=1 if i in (2, 3, 4) else 0)

    def _build_selector(self, parent, title, key, row, col):
        frame = ttk.LabelFrame(parent, text=f"{title}（勾选将被拼接进本次简历）", padding=10)
        frame.grid(row=row, column=col, sticky="nsew", padx=4, pady=4)

        canvas = tk.Canvas(frame, height=180)
        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview)
        inner = ttk.Frame(canvas)
        inner.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=inner, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        self.sel_boxes[key] = inner
        self.sel_vars[key] = []

    def save_all(self):
        self.data["education"] = self.edu_text.get("1.0", tk.END).strip()
        self.storage.save(self.data)
        self.load_selection_lists()
        messagebox.showinfo("保存成功", f"已保存到 {DATA_FILE}")

    def load_selection_lists(self):
        self.data = self.storage.load()

        self.preview_edu.delete("1.0", tk.END)
        self.preview_edu.insert("1.0", self.data.get("education", ""))

        for key, box in self.sel_boxes.items():
            for child in box.winfo_children():
                child.destroy()
            self.sel_vars[key] = []

            for item in self.data[key]:
                var = tk.BooleanVar(value=False)
                label = item.get("name") or item.get("text") or "(空)"
                cb = ttk.Checkbutton(box, text=label, variable=var)
                cb.pack(anchor="w", pady=2)
                self.sel_vars[key].append((var, item))

    def collect_selected(self):
        selected = {}
        for key, pairs in self.sel_vars.items():
            selected[key] = [item for var, item in pairs if var.get()]
        selected["education"] = self.preview_edu.get("1.0", tk.END).strip()
        selected["job"] = self.job_entry.get().strip()
        return selected

    def export(self, fmt):
        payload = self.collect_selected()

        if not payload["education"] and not any(payload[k] for k in ["internships", "projects", "skills", "evaluations"]):
            messagebox.showwarning("提示", "没有可导出的内容，请先填写并勾选。")
            return

        ext = "md" if fmt == "md" else fmt
        filename = f"resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{ext}"
        path = filedialog.asksaveasfilename(defaultextension=f".{ext}", initialfile=filename)
        if not path:
            return

        try:
            if fmt == "txt":
                self.write_txt(path, payload)
            elif fmt == "md":
                self.write_md(path, payload)
            elif fmt == "docx":
                self.write_docx(path, payload)
            elif fmt == "pdf":
                self.write_pdf(path, payload)
            else:
                raise ValueError("Unsupported format")
            messagebox.showinfo("导出成功", f"文件已导出：{path}")
        except Exception as e:
            messagebox.showerror("导出失败", str(e))

    @staticmethod
    def format_content(data):
        lines = []
        if data["job"]:
            lines.append(f"目标岗位关键词：{data['job']}")
            lines.append("")

        lines.append("# 教育经历")
        lines.append(data["education"] or "（未填写）")
        lines.append("")

        lines.append("# 实习经历")
        if data["internships"]:
            for it in data["internships"]:
                lines.append(f"- {it.get('name', '')} | {it.get('company', '')} | {it.get('period', '')}")
                lines.append(f"  职责：{it.get('duty', '')}")
        else:
            lines.append("- （未选择）")
        lines.append("")

        lines.append("# 项目经历")
        if data["projects"]:
            for it in data["projects"]:
                lines.append(f"- {it.get('name', '')} | {it.get('company', '')} | {it.get('period', '')}")
                lines.append(f"  职责：{it.get('duty', '')}")
        else:
            lines.append("- （未选择）")
        lines.append("")

        lines.append("# 个人技能")
        if data["skills"]:
            for it in data["skills"]:
                lines.append(f"- {it.get('text', '')}")
        else:
            lines.append("- （未选择）")
        lines.append("")

        lines.append("# 个人评价")
        if data["evaluations"]:
            for it in data["evaluations"]:
                lines.append(f"- {it.get('text', '')}")
        else:
            lines.append("- （未选择）")
        return "\n".join(lines)

    def write_txt(self, path, data):
        with open(path, "w", encoding="utf-8") as f:
            f.write(self.format_content(data))

    def write_md(self, path, data):
        text = self.format_content(data)
        with open(path, "w", encoding="utf-8") as f:
            f.write(text)

    def write_docx(self, path, data):
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError("未安装 python-docx，请先 pip install python-docx")

        doc = Document()
        if data["job"]:
            doc.add_paragraph(f"目标岗位关键词：{data['job']}")

        doc.add_heading("教育经历", level=1)
        doc.add_paragraph(data["education"] or "（未填写）")

        def add_exp(title, items):
            doc.add_heading(title, level=1)
            if not items:
                doc.add_paragraph("（未选择）")
                return
            for it in items:
                if "name" in it:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(f"{it.get('name','')} | {it.get('company','')} | {it.get('period','')}")
                    doc.add_paragraph(f"职责：{it.get('duty','')}")
                else:
                    doc.add_paragraph(it.get("text", ""), style="List Bullet")

        add_exp("实习经历", data["internships"])
        add_exp("项目经历", data["projects"])
        add_exp("个人技能", data["skills"])
        add_exp("个人评价", data["evaluations"])

        doc.save(path)

    def write_pdf(self, path, data):
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.pdfgen import canvas
        except ImportError:
            raise RuntimeError("未安装 reportlab，请先 pip install reportlab")

        font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
        if not os.path.exists(font_path):
            raise RuntimeError("系统缺少可用字体，无法生成 PDF")

        pdfmetrics.registerFont(TTFont("Custom", font_path))
        c = canvas.Canvas(path, pagesize=A4)
        c.setFont("Custom", 11)

        y = 800
        for line in self.format_content(data).split("\n"):
            c.drawString(40, y, line)
            y -= 18
            if y < 40:
                c.showPage()
                c.setFont("Custom", 11)
                y = 800
        c.save()


if __name__ == "__main__":
    app = ResumeGeneratorApp()
    app.mainloop()
